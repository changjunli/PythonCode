"""
创建word 模板
"""

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_LINE_SPACING
from docx.shared import Pt #磅数
from docx.oxml.ns import qn #中文格式
#以上是docx库中需要用到的部分
import time

price = input('请输入今日价格')
company_list = ['xx镇','yyy镇','zzz街']
today = time.strftime("%y{y}-%m{m}-%d{d}",time.localtime()).format(y='年',m='月',d='日')
# 获取今日时间，整理成“年-月-日”的格式

for i in company_list:
    document = Document()
    document.styles['Normal'].font.name = u'方正仿宋_GBK' #设置文档的基础字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'方正仿宋_GBK') #设置文档的基础中文字体

    p1 = document.add_paragraph() #初始化建立第一个自然段
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #设置对齐
    run1 = p1.add_run('会议通知')
    run1.font.name = '方正小标宋_GBK'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'),u'方正小标宋_GBK')
    run1.font.size = Pt(21)
    run1.font.bold = True
    p1.space_after = Pt(5)
    p1.space_before = Pt(5)

    p2 = document.add_paragraph()  # 初始化建立第一个自然段
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run2 = p2.add_run(i + ':') #这里是对客户的称呼
    run2.font.name = '方正仿宋_GBK'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'方正仿宋_GBK')
    run2.font.size = Pt(16)
    run2.font.bold = True

    p3 = document.add_paragraph()  # 初始化建立第一个自然段
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST  # 1.5倍行距
    run3 = p3.add_run('    根据公司安排，为提供优质客户，拟定了今日黄金价格为%s，特此通知。'%price)  # 这里是对客户的称呼
    run3.font.name = '方正仿宋_GBK'
    print(p3.paragraph_format.line_spacing_rule,p3.paragraph_format.line_spacing)
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'方正仿宋_GBK')
    run3.font.size = Pt(16)
    run3.font.bold = True

    p4 = document.add_paragraph()  # 初始化建立第一个自然段
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run4 = p4.add_run('xxxxxx室')  # 这里是对客户的称呼
    run4.font.name = '方正仿宋_GBK'
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'方正仿宋_GBK')
    run4.font.size = Pt(16)
    run4.font.bold = True
    p1.space_before = Pt(25)

    p5 = document.add_paragraph()  # 初始化建立第一个自然段
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run5 = p5.add_run('(联系人：小刘  联系电话：18888888888)')  # 这里是对客户的称呼
    run5.font.name = '方正仿宋_GBK'
    run5.element.rPr.rFonts.set(qn('w:eastAsia'), u'方正仿宋_GBK')
    run5.font.size = Pt(16)
    run5.font.bold = True


    document.save('发%s会议通知.docx'%i)