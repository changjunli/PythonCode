from docx import Document

doc = Document('案例.docx')

for i in range(len(doc.paragraphs)):
        print(doc.paragraphs[i].text)

tabs = doc.tables
for tb in tabs:
        for row in tb.rows:
                 for cell in row.cells:
                #         print(cell.text)
                        text = ''
                        for p in cell.pargraphs:
                                text += p.text
                                print(text)
